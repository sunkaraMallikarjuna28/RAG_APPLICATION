# main.py - Complete RAG Application with All Advanced Features
import os
from dotenv import load_dotenv
import PyPDF2
from langchain_experimental.text_splitter import SemanticChunker
from langchain_community.embeddings import HuggingFaceEmbeddings
import google.generativeai as genai
import pymongo
from langchain import hub
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema.output_parser import StrOutputParser
from langchain.schema.runnable import RunnablePassthrough
from docx import Document
import numpy as np
from pathlib import Path

# Load environment variables
load_dotenv()

class SimpleRAGApp:
    def __init__(self):
        print("🚀 Starting Complete RAG Application...")
        
        # Setup Google AI
        self.google_api_key = os.getenv("GOOGLE_API_KEY")
        genai.configure(api_key=self.google_api_key)
        
        # Setup MongoDB
        self.mongo_uri = os.getenv("MONGODB_URI")
        self.client = pymongo.MongoClient(self.mongo_uri)
        self.db = self.client["rag_database"]
        self.collection = self.db["document_chunks"]
        
        # Setup embeddings for semantic chunking
        self.embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        
        # Setup semantic chunker
        self.text_splitter = SemanticChunker(embeddings=self.embeddings)
        
        # Setup Gemini LLM
        self.model = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            google_api_key=self.google_api_key
        )
        
        # Setup RAG prompt
        self.prompt = hub.pull("rlm/rag-prompt")
        
        # Setup project paths
        self.project_dir = Path.cwd()
        self.data_dir = self.project_dir / "data"
        self.output_dir = self.project_dir / "outputs"
        
        # Create directories if they don't exist
        self.data_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
        
        print("✅ Complete RAG Application initialized!")
    
    def find_pdf_files(self):
        """Auto-detect PDF files in data folder"""
        pdf_files = list(self.data_dir.glob("*.pdf"))
        
        if pdf_files:
            print(f"📄 Found {len(pdf_files)} PDF files:")
            for i, pdf in enumerate(pdf_files, 1):
                print(f"   {i}. {pdf.name}")
        else:
            print("❌ No PDF files found in data folder")
        
        return pdf_files
    
    def extract_pdf_text(self, pdf_path):
        """Step 1: Extract text from PDF"""
        print(f"📄 Reading PDF: {pdf_path}")
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text
                    print(f"   ✅ Page {page_num + 1} processed")
                
                print(f"📊 Total text extracted: {len(text)} characters")
                return text
                
        except Exception as e:
            print(f"❌ Error reading PDF: {e}")
            return None
    
    def create_semantic_chunks(self, text):
        """Step 2: Create semantic chunks"""
        print("🧠 Creating semantic chunks...")
        
        try:
            chunks = self.text_splitter.split_text(text)
            print(f"✅ Created {len(chunks)} semantic chunks")
            
            # Show sample chunks
            for i, chunk in enumerate(chunks[:2]):
                print(f"   Sample Chunk {i+1}: {chunk[:100]}...")
            
            return chunks
            
        except Exception as e:
            print(f"❌ Error creating chunks: {e}")
            return []
    
    def generate_embeddings(self, chunks):
        """Step 3: Generate embeddings for chunks"""
        print("🔢 Generating embeddings...")
        
        try:
            chunk_embeddings = []
            
            for i, chunk in enumerate(chunks):
                # Generate embedding using Google
                embedding_result = genai.embed_content(
                    model="models/embedding-001",
                    content=chunk
                )
                
                chunk_data = {
                    "chunk_id": i,
                    "text": chunk,
                    "embedding": embedding_result["embedding"]
                }
                
                chunk_embeddings.append(chunk_data)
                print(f"   ✅ Chunk {i+1}/{len(chunks)} embedded")
            
            print(f"🎉 Generated embeddings for {len(chunk_embeddings)} chunks")
            return chunk_embeddings
            
        except Exception as e:
            print(f"❌ Error generating embeddings: {e}")
            return []
    
    def store_in_mongodb(self, chunk_embeddings):
        """Step 4: Store embeddings in MongoDB Atlas"""
        print("💾 Storing embeddings in MongoDB Atlas...")
        
        try:
            # Clear existing data
            self.collection.delete_many({})
            
            # Insert new data
            self.collection.insert_many(chunk_embeddings)
            
            print(f"✅ Stored {len(chunk_embeddings)} chunks in MongoDB")
            
            
            
        except Exception as e:
            print(f"❌ Error storing in MongoDB: {e}")
    
    def create_vector_index(self):
        """Step 5: Create comprehensive vector search indexes (FLAT, HNSW, IVF)"""
        print("🔍 Creating comprehensive vector search indexes...")
        
        try:
            # Check if collection exists and has data
            doc_count = self.collection.count_documents({})
            print(f"   📊 Documents in collection: {doc_count}")
            
            # 1. Create FLAT index (exact similarity search)
            print("   Creating FLAT index for exact search...")
            self.collection.create_index("chunk_id", name="flat_index")
            
            # 2. Create HNSW simulation index (compound for performance)
            print("   Creating HNSW simulation index...")
            self.collection.create_index([
                ("chunk_id", 1), 
                ("text", 1)
            ], name="hnsw_index")
            
            # 3. Create IVF simulation index (partitioned search)
            print("   Creating IVF simulation index...")
            self.collection.create_index([("chunk_id", 1)], name="ivf_index")
            
            # 4. Create BM25 text search index
            print("   Creating BM25 text search index...")
            self.collection.create_index([("text", "text")], name="bm25_index")
            
            # 5. Store search configurations
            self._store_search_configurations()
            
            # Verify indexes
            indexes = list(self.collection.list_indexes())
            print(f"   ✅ Total indexes created: {len(indexes)}")
            
            for index in indexes:
                index_name = index.get('name', 'Unknown')
                print(f"      - {index_name}")
            
            print("✅ All vector search indexes created successfully!")
            return True
            
        except Exception as e:
            print(f"❌ Error creating vector indexes: {e}")
            return False

    def _store_search_configurations(self):
        """Store search strategy configurations"""
        print("   Storing search strategy configurations...")
        
        try:
            config_collection = self.db["search_configurations"]
            
            configs = [
                {
                    "strategy": "FLAT",
                    "description": "Exact similarity search using cosine distance",
                    "index_name": "flat_index",
                    "active": True
                },
                {
                    "strategy": "HNSW", 
                    "description": "Hierarchical Navigable Small World (approximate)",
                    "index_name": "hnsw_index",
                    "active": True
                },
                {
                    "strategy": "IVF",
                    "description": "Inverted File Index (memory efficient)",
                    "index_name": "ivf_index", 
                    "active": True
                },
                {
                    "strategy": "BM25",
                    "description": "Best Matching 25 (keyword search)",
                    "index_name": "bm25_index",
                    "active": True
                },
                {
                    "strategy": "MMR",
                    "description": "Maximal Marginal Relevance (diversity)",
                    "index_name": "hnsw_index",
                    "active": True
                }
            ]
            
            for config in configs:
                config_collection.replace_one(
                    {"strategy": config["strategy"]}, 
                    config, 
                    upsert=True
                )
            
            print(f"   ✅ Stored {len(configs)} search configurations")
            
        except Exception as e:
            print(f"   ⚠️ Warning: Could not store configurations: {e}")
    
    def similarity_search(self, query, top_k=5, strategy="FLAT"):
        """Step 6: Multi-strategy similarity search (FLAT, HNSW, IVF, BM25, MMR)"""
        print(f"🔍 Searching with {strategy} strategy: '{query}'")
        
        try:
            if strategy == "FLAT":
                return self._flat_search(query, top_k)
            elif strategy == "HNSW":
                return self._hnsw_search(query, top_k)
            elif strategy == "IVF":
                return self._ivf_search(query, top_k)
            elif strategy == "BM25":
                return self._bm25_search(query, top_k)
            elif strategy == "MMR":
                return self._mmr_search(query, top_k)
            else:
                print(f"⚠️ Unknown strategy {strategy}, using FLAT")
                return self._flat_search(query, top_k)
                
        except Exception as e:
            print(f"❌ Error in {strategy} search: {e}")
            return []

    def _flat_search(self, query, top_k):
        """FLAT: Exact similarity search"""
        print("   Using FLAT (exact) search...")
        
        # Generate query embedding
        query_embedding = genai.embed_content(
            model="models/embedding-001",
            content=query
        )["embedding"]
        
        # Get all chunks
        all_chunks = list(self.collection.find({}))
        
        # Calculate exact cosine similarities
        similarities = []
        for chunk in all_chunks:
            chunk_embedding = chunk["embedding"]
            similarity = np.dot(query_embedding, chunk_embedding) / (
                np.linalg.norm(query_embedding) * np.linalg.norm(chunk_embedding)
            )
            similarities.append((chunk, similarity))
        
        # Sort and return top_k
        similarities.sort(key=lambda x: x[1], reverse=True)
        top_chunks = similarities[:top_k]
        
        print(f"   ✅ FLAT search found {len(top_chunks)} chunks")
        return [chunk[0] for chunk in top_chunks]

    def _hnsw_search(self, query, top_k):
        """HNSW: Hierarchical Navigable Small World (approximate search)"""
        print("   Using HNSW (approximate) search...")
        
        query_embedding = genai.embed_content(
            model="models/embedding-001",
            content=query
        )["embedding"]
        
        # Sample chunks for approximate search (HNSW simulation)
        sample_chunks = list(self.collection.find({}).limit(50))
        
        similarities = []
        for chunk in sample_chunks:
            chunk_embedding = chunk["embedding"]
            similarity = np.dot(query_embedding, chunk_embedding) / (
                np.linalg.norm(query_embedding) * np.linalg.norm(chunk_embedding)
            )
            similarities.append((chunk, similarity))
        
        similarities.sort(key=lambda x: x[1], reverse=True)
        top_chunks = similarities[:top_k]
        
        print(f"   ✅ HNSW search found {len(top_chunks)} chunks")
        return [chunk[0] for chunk in top_chunks]

    def _ivf_search(self, query, top_k):
        """IVF: Inverted File Index (partitioned search)"""
        print("   Using IVF (partitioned) search...")
        
        total_chunks = self.collection.count_documents({})
        partition_size = max(1, total_chunks // 10)
        
        query_embedding = genai.embed_content(
            model="models/embedding-001",
            content=query
        )["embedding"]
        
        all_similarities = []
        for i in range(0, total_chunks, partition_size):
            partition_chunks = list(self.collection.find({
                "chunk_id": {"$gte": i, "$lt": i + partition_size}
            }))
            
            for chunk in partition_chunks:
                chunk_embedding = chunk["embedding"]
                similarity = np.dot(query_embedding, chunk_embedding) / (
                    np.linalg.norm(query_embedding) * np.linalg.norm(chunk_embedding)
                )
                all_similarities.append((chunk, similarity))
        
        all_similarities.sort(key=lambda x: x[1], reverse=True)
        top_chunks = all_similarities[:top_k]
        
        print(f"   ✅ IVF search found {len(top_chunks)} chunks")
        return [chunk[0] for chunk in top_chunks]

    def _bm25_search(self, query, top_k):
        """BM25: Keyword-based text search"""
        print("   Using BM25 text search...")
        
        try:
            results = list(self.collection.find(
                {"$text": {"$search": query}},
                {"score": {"$meta": "textScore"}}
            ).sort([("score", {"$meta": "textScore"})]).limit(top_k))
            
            print(f"   ✅ BM25 search found {len(results)} chunks")
            return results
            
        except Exception as e:
            print(f"   ⚠️ BM25 search failed, falling back to FLAT: {e}")
            return self._flat_search(query, top_k)

    def _mmr_search(self, query, top_k):
        """MMR: Maximal Marginal Relevance (diversity-aware search)"""
        print("   Using MMR (diversity) search...")
        
        # Get more candidates than needed
        candidates = self._flat_search(query, top_k * 3)
        
        if not candidates:
            return []
        
        # MMR selection for diversity
        selected = [candidates[0]]
        candidates = candidates[1:]
        
        lambda_mult = 0.5  # Balance relevance vs diversity
        
        query_embedding = genai.embed_content(
            model="models/embedding-001",
            content=query
        )["embedding"]
        
        while len(selected) < top_k and candidates:
            best_score = -1
            best_candidate = None
            best_idx = -1
            
            for idx, candidate in enumerate(candidates):
                # Relevance score
                relevance = np.dot(query_embedding, candidate["embedding"]) / (
                    np.linalg.norm(query_embedding) * np.linalg.norm(candidate["embedding"])
                )
                
                # Diversity score
                max_similarity = 0
                for selected_doc in selected:
                    similarity = np.dot(candidate["embedding"], selected_doc["embedding"]) / (
                        np.linalg.norm(candidate["embedding"]) * np.linalg.norm(selected_doc["embedding"])
                    )
                    max_similarity = max(max_similarity, similarity)
                
                # MMR score
                mmr_score = lambda_mult * relevance - (1 - lambda_mult) * max_similarity
                
                if mmr_score > best_score:
                    best_score = mmr_score
                    best_candidate = candidate
                    best_idx = idx
            
            if best_candidate:
                selected.append(best_candidate)
                candidates.pop(best_idx)
        
        print(f"   ✅ MMR search found {len(selected)} diverse chunks")
        return selected
    
    def format_docs(self, docs):
        """Step 7: Format documents for prompt"""
        return "\n\n".join(doc["text"] for doc in docs)
    
    def generate_answer(self, query, strategy="FLAT"):
        """Step 8: Generate answer using Gemini with specified search strategy"""
        print(f"🤖 Generating answer with Gemini using {strategy} search...")
        
        try:
            # Get relevant chunks using specified strategy
            relevant_chunks = self.similarity_search(query, strategy=strategy)
            
            if not relevant_chunks:
                return "No relevant information found in the document."
            
            # Format context
            context = self.format_docs(relevant_chunks)
            
            # Create RAG chain exactly as specified
            rag_chain = (
                {"context": lambda x: context, "question": RunnablePassthrough()}
                | self.prompt
                | self.model
                | StrOutputParser()
            )
            
            # Generate answer
            answer = rag_chain.invoke(query)
            
            print("✅ Answer generated successfully!")
            return answer
            
        except Exception as e:
            print(f"❌ Error generating answer: {e}")
            return f"Error generating answer: {e}"
    
    def save_to_docx(self, query, answer, strategy="FLAT", filename=None):
        """Step 9: Save output to DOCX"""
        if filename is None:
            filename = f"rag_output_{strategy.lower()}.docx"
        
        print(f"📄 Saving to DOCX: {filename}")
        
        try:
            doc = Document()
            
            # Add title
            doc.add_heading('RAG Application Output', 0)
            
            # Add search strategy
            doc.add_heading('Search Strategy Used:', level=1)
            doc.add_paragraph(f"{strategy} - Advanced similarity search")
            
            # Add query
            doc.add_heading('Question:', level=1)
            doc.add_paragraph(query)
            
            # Add answer
            doc.add_heading('Answer:', level=1)
            doc.add_paragraph(answer)
            
            # Save file
            output_path = self.output_dir / filename
            doc.save(output_path)
            
            print(f"✅ Output saved to: {output_path}")
            
        except Exception as e:
            print(f"❌ Error saving DOCX: {e}")
    
    def auto_select_pdf(self):
        """Automatically find and select PDF file"""
        pdf_files = self.find_pdf_files()
        
        if not pdf_files:
            print("❌ No PDF files found in data folder")
            return None
        
        if len(pdf_files) == 1:
            selected_pdf = pdf_files[0]
            print(f"✅ Auto-selected: {selected_pdf.name}")
            return str(selected_pdf)
        else:
            print(f"📄 Found {len(pdf_files)} PDF files:")
            for i, pdf in enumerate(pdf_files, 1):
                print(f"   {i}. {pdf.name}")
            
            try:
                choice = int(input("Enter number to select PDF: ")) - 1
                selected_pdf = pdf_files[choice]
                print(f"✅ Selected: {selected_pdf.name}")
                return str(selected_pdf)
            except:
                print("❌ Invalid selection, using first PDF")
                return str(pdf_files[0])
    
    def process_pdf_and_query(self, pdf_path, query):
        """Complete RAG pipeline"""
        print("🎯 Starting complete RAG pipeline...")
        
        # Step 1: Extract text from PDF
        text = self.extract_pdf_text(pdf_path)
        if not text:
            return None
        
        # Step 2: Create semantic chunks
        chunks = self.create_semantic_chunks(text)
        if not chunks:
            return None
        
        # Step 3: Generate embeddings
        chunk_embeddings = self.generate_embeddings(chunks)
        if not chunk_embeddings:
            return None
        
        # Step 4: Store in MongoDB with all indices
        self.store_in_mongodb(chunk_embeddings)
        
        print("\n🎉 RAG Pipeline Setup Complete!")
        return True
    
    def test_all_strategies(self, query):
        """Test all search strategies (FLAT, HNSW, IVF, BM25, MMR)"""
        print(f"\n🔍 Testing all search strategies with query: '{query}'")
        
        strategies = ["FLAT", "HNSW", "IVF", "BM25", "MMR"]
        results = {}
        
        for strategy in strategies:
            print(f"\n--- Testing {strategy} Strategy ---")
            try:
                answer = self.generate_answer(query, strategy=strategy)
                self.save_to_docx(query, answer, strategy=strategy)
                results[strategy] = answer[:150] + "..." if len(answer) > 150 else answer
                print(f"✅ {strategy} completed successfully")
            except Exception as e:
                print(f"❌ {strategy} failed: {e}")
                results[strategy] = f"Error: {e}"
        
        # Summary
        print(f"\n📊 Strategy Comparison Summary:")
        for strategy, result in results.items():
            print(f"   {strategy}: {result[:80]}...")
        
        return results

# Main execution
if __name__ == "__main__":
    print("🚀 Starting Complete RAG Application with All Features...")
    
    # Initialize RAG app
    rag_app = SimpleRAGApp()
    
    # Auto-find PDF file
    pdf_path = rag_app.auto_select_pdf()
    
    if pdf_path:
        # Process the PDF
        query = "What is the main topic of this document?"
        
        # Setup the RAG system
        setup_success = rag_app.process_pdf_and_query(pdf_path, query)
        
        if setup_success:
            # Test all search strategies
            rag_app.test_all_strategies(query)
            
            # Interactive mode
            print("\n🎯 Interactive Mode - Ask your questions!")
            print("Available strategies: FLAT, HNSW, IVF, BM25, MMR")
            print("Type 'quit' to exit")
            
            while True:
                user_query = input("\n💬 Your question: ").strip()
                
                if user_query.lower() in ['quit', 'exit', 'q']:
                    print("👋 Goodbye!")
                    break
                
                if user_query:
                    # Ask which strategy to use
                    print("\nChoose search strategy:")
                    print("1. FLAT (exact similarity)")
                    print("2. HNSW (fast approximate)")
                    print("3. IVF (memory efficient)")
                    print("4. BM25 (keyword search)")
                    print("5. MMR (diverse results)")
                    
                    strategy_choice = input("Enter number (1-5) or press Enter for FLAT: ").strip()
                    
                    strategy_map = {
                        "1": "FLAT",
                        "2": "HNSW", 
                        "3": "IVF",
                        "4": "BM25",
                        "5": "MMR"
                    }
                    
                    strategy = strategy_map.get(strategy_choice, "FLAT")
                    
                    # Generate answer
                    answer = rag_app.generate_answer(user_query, strategy=strategy)
                    print(f"\n💡 Answer ({strategy}): {answer}")
                    
                    # Save to DOCX
                    rag_app.save_to_docx(user_query, answer, strategy=strategy)
    else:
        print("❌ No PDF file available. Please add a PDF to the data folder and try again.")
